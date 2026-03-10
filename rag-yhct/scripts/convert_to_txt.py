#!/usr/bin/env python3
"""
Convert all PDF and DOCX files in data/raw/ to plain UTF-8 text files in raw-txt/.

Usage:
    python scripts/convert_to_txt.py [--raw-dir data/raw] [--out-dir raw-txt]

The output directory mirrors the source directory structure.
Files that have already been converted are skipped unless --force is passed.

Supported formats  : .pdf, .docx, .csv
Skipped silently   : .Zone.Identifier, .gitkeep, directories
"""

from __future__ import annotations

import argparse
import logging
import sys
from pathlib import Path
from typing import Callable

# ---------------------------------------------------------------------------
# TCVN3 (.VnTime) → Unicode conversion
# ---------------------------------------------------------------------------
# .VnTime fonts place Vietnamese glyphs at Latin-1 / CP1252 code positions.
# PyMuPDF extracts the raw byte values and maps them through the font's
# ToUnicode CMap, producing "garbled" Latin-1 characters.  This table
# converts those garbled characters back to correct Vietnamese Unicode.
# Verified against 43 word samples from actual garbled PDFs.

_VNTIME_TO_UNICODE: dict[str, str] = {
    # ---- vowel base + diacritics (lowercase) ----
    '\u00A8': 'ă',  '\u00A9': 'â',  '\u00AA': 'ê',  '\u00AB': 'ô',
    '\u00AC': 'ơ',  '\u00AD': 'ư',  '\u00AE': 'đ',
    '\u00B5': 'à',  '\u00B6': 'ả',  '\u00B7': 'ã',  '\u00B8': 'á',
    '\u00B9': 'ạ',  '\u00BB': 'ằ',  '\u00BE': 'ắ',
    '\u00C6': 'ặ',  '\u00C7': 'ầ',  '\u00C8': 'ẩ',  '\u00C9': 'ẫ',
    '\u00CA': 'ấ',  '\u00CB': 'ậ',  '\u00CC': 'ì',  '\u00CE': 'ẻ',
    '\u00CF': 'ĩ',
    '\u00D0': 'é',  '\u00D1': 'ẹ',  '\u00D2': 'ề',  '\u00D3': 'ể',
    '\u00D4': 'ễ',  '\u00D5': 'ế',  '\u00D6': 'ệ',  '\u00D7': 'ì',
    '\u00D8': 'ỉ',  '\u00DC': 'ĩ',  '\u00DD': 'í',  '\u00DE': 'ị',
    '\u00DF': 'ò',
    '\u00E1': 'ỏ',  '\u00E2': 'õ',  '\u00E3': 'ó',  '\u00E4': 'ọ',
    '\u00E5': 'ồ',  '\u00E6': 'ổ',  '\u00E7': 'ỗ',  '\u00E8': 'ố',
    '\u00E9': 'ộ',  '\u00EA': 'ờ',  '\u00EB': 'ở',  '\u00ED': 'ớ',
    '\u00EE': 'ợ',  '\u00EF': 'ù',
    '\u00F1': 'ủ',  '\u00F2': 'ũ',  '\u00F3': 'ú',  '\u00F4': 'ụ',
    '\u00F5': 'ừ',  '\u00F6': 'ử',  '\u00F7': 'ữ',  '\u00F8': 'ứ',
    '\u00F9': 'ự',  '\u00FA': 'ỵ',  '\u00FB': 'ỷ',  '\u00FD': 'ý',
    '\u00FE': 'ỹ',
    # ---- uppercase Đ ----
    '\u00A7': 'Đ',
}

_VNTIME_TRANS = str.maketrans(_VNTIME_TO_UNICODE)

# Markers used to detect TCVN3-garbled text
_TCVN3_MARKERS = frozenset(chr(cp) for cp in [
    0xA7, 0xA8, 0xAB, 0xAC, 0xAE,
    0xB5, 0xB6, 0xB7, 0xB8, 0xB9, 0xBB, 0xBE,
    0xC6, 0xC7, 0xC8, 0xC9, 0xCA, 0xCB, 0xCC, 0xCE, 0xCF,
    0xD0, 0xD1, 0xD2, 0xD3, 0xD4, 0xD5, 0xD6, 0xD7, 0xD8,
    0xDC, 0xDD, 0xDE, 0xDF,
    0xE4, 0xE5, 0xE6, 0xE7, 0xE8, 0xE9, 0xEA, 0xEB, 0xEE, 0xEF,
    0xF1, 0xF5, 0xF6, 0xF7, 0xF8, 0xF9, 0xFA, 0xFB, 0xFE,
])
_TCVN3_STRONG = frozenset(chr(cp) for cp in [0xAE, 0xB8, 0xB5])


def _is_tcvn3_garbled(text: str) -> bool:
    """Return True if *text* looks like TCVN3 (.VnTime) garbled output."""
    if len(text) < 100:
        return False
    marker_count = sum(1 for ch in text if ch in _TCVN3_MARKERS)
    strong_count = sum(1 for ch in text if ch in _TCVN3_STRONG)
    ratio = marker_count / len(text)
    return ratio > 0.02 or strong_count > 10


def _tcvn3_to_unicode(text: str) -> str:
    """Convert TCVN3-garbled text to proper Vietnamese Unicode."""
    return text.translate(_VNTIME_TRANS)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PDF  →  text
# ---------------------------------------------------------------------------

def _pdf_to_text(pdf_path: Path) -> str:
    """Extract all text from a PDF using PyMuPDF (preserves Unicode / tiếng Việt).

    Automatically detects and converts TCVN3 (.VnTime) garbled text to Unicode.
    """
    import fitz  # pyright: ignore[reportMissingTypeStubs]  # PyMuPDF

    pages: list[str] = []
    try:
        doc = fitz.open(str(pdf_path))
    except Exception as exc:
        raise RuntimeError(f"Cannot open PDF: {exc}") from exc

    for page in doc:
        text: str = page.get_text("text")  # type: ignore[attr-defined]
        pages.append(text)

    doc.close()
    full_text = "\n\f\n".join(pages)   # \f = form-feed, one section per page

    if _is_tcvn3_garbled(full_text):
        logger.info("  → TCVN3 detected, converting to Unicode: %s", pdf_path.name)
        full_text = _tcvn3_to_unicode(full_text)

    return full_text


# ---------------------------------------------------------------------------
# DOCX  →  text
# ---------------------------------------------------------------------------

def _docx_to_text(docx_path: Path) -> str:
    """Extract all text from a DOCX preserving paragraph order."""
    import docx  # python-docx

    try:
        doc = docx.Document(str(docx_path))
    except Exception as exc:
        raise RuntimeError(f"Cannot open DOCX: {exc}") from exc

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                paragraphs.append("\t".join(row_cells))

    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# CSV  →  text (TSV: one row per line, cells tab-separated)
# ---------------------------------------------------------------------------

def _csv_to_text(csv_path: Path) -> str:
    """Convert a CSV to tab-separated plain text.

    Encoding detection order: UTF-8-BOM → UTF-8 → cp1252 (Windows Vietnamese).
    Each data row becomes one line; cells are joined with a tab character.
    """
    import csv

    encodings = ("utf-8-sig", "utf-8", "cp1252")
    raw_bytes = csv_path.read_bytes()

    for enc in encodings:
        try:
            content = raw_bytes.decode(enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    else:
        # Last resort: decode with replacement
        content = raw_bytes.decode("utf-8", errors="replace")

    rows: list[str] = []
    for row in csv.reader(content.splitlines()):
        cells = [cell.strip() for cell in row]
        line = "\t".join(cells)
        if line.strip():
            rows.append(line)

    return "\n".join(rows)


# ---------------------------------------------------------------------------
# Core conversion
# ---------------------------------------------------------------------------

CONVERTERS: dict[str, Callable[[Path], str]] = {
    ".pdf": _pdf_to_text,
    ".docx": _docx_to_text,
    ".csv": _csv_to_text,
}


def convert_file(src: Path, dst: Path) -> bool:
    """Convert *src* to UTF-8 text and write to *dst*.

    Returns True on success, False on failure.
    """
    ext = src.suffix.lower()
    converter = CONVERTERS.get(ext)
    if converter is None:
        logger.debug("Skipped (unsupported type): %s", src)
        return False

    # Skip Excel binary files (.xls old format) misnamed as .csv
    if ext == ".csv":
        try:
            header = src.read_bytes()[:8]
            # OLE2/Composite Document magic: D0 CF 11 E0 A1 B1 1A E1
            if header[:8] == b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1':
                logger.warning("Skipped (Excel binary, not CSV): %s", src)
                return False
        except Exception:
            pass  # if read fails, let converter handle it

    try:
        text = converter(src)
    except Exception as exc:
        logger.error("FAILED %s — %s", src, exc)
        return False

    if not text.strip():
        logger.warning("Empty text extracted from %s — skipping", src)
        return False

    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(text, encoding="utf-8")
    return True


def convert_directory(raw_dir: Path, out_dir: Path, force: bool = False) -> None:
    """Walk *raw_dir* and convert every supported file into *out_dir*."""
    supported_exts = set(CONVERTERS)
    files = sorted(
        p for p in raw_dir.rglob("*")
        if p.is_file() and p.suffix.lower() in supported_exts
    )

    if not files:
        logger.warning("No supported files found in %s", raw_dir)
        return

    logger.info("Found %d file(s) to process in %s", len(files), raw_dir)
    ok = fail = skip = 0

    for src in files:
        rel = src.relative_to(raw_dir)
        dst = out_dir / rel.with_suffix(".txt")

        if dst.exists() and not force:
            logger.debug("Already exists, skipping: %s", dst)
            skip += 1
            continue

        logger.info("Converting: %s", rel)
        if convert_file(src, dst):
            ok += 1
        else:
            fail += 1

    logger.info(
        "Done — converted: %d | failed: %d | skipped (already exists): %d",
        ok, fail, skip,
    )


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    here = Path(__file__).resolve().parent.parent   # project root (rag-yhct/)

    parser = argparse.ArgumentParser(
        description="Convert PDF/DOCX files to UTF-8 TXT (giữ nguyên font tiếng Việt)"
    )
    parser.add_argument(
        "--raw-dir",
        type=Path,
        default=here / "data" / "raw",
        help="Thư mục chứa file gốc (mặc định: data/raw)",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=here / "data" / "raw-txt",
        help="Thư mục xuất file TXT (mặc định: data/raw-txt/)",
    )
    parser.add_argument(
        "--force",
        action="store_true",
        help="Ghi đè file TXT đã tồn tại",
    )
    args = parser.parse_args()

    raw_dir: Path = args.raw_dir.resolve()
    out_dir: Path = args.out_dir.resolve()

    if not raw_dir.exists():
        logger.error("raw-dir không tồn tại: %s", raw_dir)
        sys.exit(1)

    logger.info("Nguồn  : %s", raw_dir)
    logger.info("Đích   : %s", out_dir)

    convert_directory(raw_dir, out_dir, force=args.force)


if __name__ == "__main__":
    main()
